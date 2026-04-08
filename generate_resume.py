"""生成 HoyoCalendar 项目简历描述 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.25

# ======== 项目标题 ========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run('HoyoCalendar — 智能桌面日历应用')
run.bold = True
run.font.size = Pt(14)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ======== 项目描述 ========
desc_p = doc.add_paragraph()
label = desc_p.add_run('项目描述：')
label.bold = True
label.font.size = Pt(11)
label.font.name = '宋体'
label._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
text = desc_p.add_run(
    '本项目是一个基于 Electron + Flask 的桌面悬浮日历应用，采用二次元风格玻璃拟态 UI。'
    '后端提供日程 CRUD RESTful API，支持一次性日程和长期循环任务（每天/每周/每月），'
    '并集成多家 AI 大模型（豆包/Ollama/OpenAI），实现自然语言解析和基于 Function Call 的对话式日程管理功能。'
)
text.font.size = Pt(11)
text.font.name = '宋体'
text._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ======== 使用技术 ========
tech_p = doc.add_paragraph()
label2 = tech_p.add_run('使用技术：')
label2.bold = True
label2.font.size = Pt(11)
label2.font.name = '宋体'
label2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
tech = tech_p.add_run('Python、Flask、Flask-CORS、OpenAI SDK、RESTful API、Function Call、JSON、PyInstaller')
tech.font.size = Pt(11)
tech.font.name = '宋体'
tech._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ======== 工作职责 ========
duty_label = doc.add_paragraph()
dl_run = duty_label.add_run('工作职责：')
dl_run.bold = True
dl_run.font.size = Pt(11)
dl_run.font.name = '宋体'
dl_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

duties = [
    [
        ('基于 Flask 搭建后端 RESTful API 服务，实现日程的', False),
        ('CRUD 操作', True),
        ('与', False),
        ('JSON 文件持久化', True),
        ('，支持按日期过滤与关键词模糊检索；设计循环日程引擎，涵盖 daily/weekly/monthly 三种', False),
        ('周期调度策略', True),
        ('，内置日期范围判定、打卡状态追踪及完成度计算逻辑；', False),
    ],
    [
        ('通过 OpenAI SDK 封装统一的', False),
        ('多 Provider 适配层', True),
        ('（豆包/Ollama/OpenAI），基于', False),
        ('Function Call', True),
        ('机制驱动 AI 自主调用工具函数完成日程操作，支持最多 6 轮', False),
        ('Tool-Use 循环', True),
        ('；同时实现', False),
        ('自然语言解析', True),
        ('与', False),
        ('多模态图片识别', True),
        ('接口，将非结构化输入转为标准日程 JSON；', False),
    ],
    [
        ('采用', False),
        ('时间戳 + 原子计数器 + 互斥锁', True),
        ('的线程安全 ID 生成方案保证并发唯一性；通过 Flask-CORS 处理跨域请求；使用 PyInstaller 将后端打包为', False),
        ('独立可执行文件', True),
        ('，配合 Electron 主进程实现', False),
        ('动态端口探测', True),
        ('与子进程生命周期管理，支持一键分发部署。', False),
    ],
]

for i, parts in enumerate(duties, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    # 序号
    num_run = p.add_run(f'{i}. ')
    num_run.font.size = Pt(11)
    num_run.font.name = '宋体'
    num_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for text_content, is_bold in parts:
        run = p.add_run(text_content)
        run.bold = is_bold
        run.font.size = Pt(11)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

output_path = r'd:\vscodeprojects\hoyoCalendar\HoyoCalendar_简历项目描述.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
